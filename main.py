from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import qrcode
from PIL import Image
import os
from typing import List
from datetime import datetime
import io
import shutil

app = FastAPI(title="Lab Record Generator API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Experiment(BaseModel):
    title: str
    date: str = ""
    github: str

class RecordData(BaseModel):
    course_title: str
    student_name: str
    register_number: str
    experiments: List[Experiment]

def create_qr_code(url: str, size: int = 200):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    img = img.resize((size, size), Image.Resampling.LANCZOS)

    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    return img_byte_arr

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'tc{}'.format(edge.capitalize())
            element = OxmlElement('w:{}'.format(tag))
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)

@app.get("/")
async def root():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logo_exists = any(
        os.path.exists(os.path.join(base_dir, f"college_logo.{ext}"))
        for ext in ["png", "jpg", "jpeg"]
    )
    return {
        "message": "Lab Record Generator API",
        "status": "running",
        "version": "1.0",
        "logo_uploaded": logo_exists
    }

@app.post("/upload-logo")
async def upload_logo(file: UploadFile = File(...)):
    try:
        if not file.content_type.startswith('image/'):
            raise HTTPException(status_code=400, detail="File must be an image")
        
        ext = file.filename.split('.')[-1].lower()
        if ext not in ['png', 'jpg', 'jpeg']:
            raise HTTPException(status_code=400, detail="Only PNG, JPG, JPEG files allowed")
        
        base_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(base_dir, f"college_logo.{ext}")

        with open(logo_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        if ext in ['jpg', 'jpeg']:
            img = Image.open(logo_path)
            png_path = os.path.join(base_dir, "college_logo.png")
            img.save(png_path)
            os.remove(logo_path)
            logo_path = png_path
        
        return {"message": "Logo uploaded successfully", "filename": logo_path}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading logo: {str(e)}")

@app.post("/generate-docx")
async def generate_docx(data: RecordData):
    try:
        doc = Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # ✅ FIXED LOGO PATH FOR RENDER
        base_dir = os.path.dirname(os.path.abspath(__file__))
        logo_files = ['college_logo.png', 'college_logo.jpg', 'college_logo.jpeg']
        logo_exists = False

        for name in logo_files:
            logo_path = os.path.join(base_dir, name)
            if os.path.exists(logo_path):
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(logo_path, width=Inches(7.0))
                logo_exists = True
                break

        if logo_exists:
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(data.course_title)
        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph()

        num_experiments = len(data.experiments)
        table = doc.add_table(rows=num_experiments + 1, cols=6)
        table.style = 'Table Grid'

        headers = ['Exp', 'Date', 'Name of The Experiment', 'QR Code', 'Mark', 'Signature']
        header_cells = table.rows[0].cells
        
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cell, top=1, bottom=1, left=1, right=1)

        widths = [Inches(0.5), Inches(0.8), Inches(3.0), Inches(0.8), Inches(0.6), Inches(1.0)]
        for idx, width in enumerate(widths):
            for row in table.rows:
                row.cells[idx].width = width

        qr_images = []

        for idx, exp in enumerate(data.experiments):
            row = table.rows[idx + 1]
            cells = row.cells

            cells[0].text = str(idx + 1).zfill(2)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            cells[1].text = exp.date if exp.date else ""

            title_para = cells[2].paragraphs[0]
            title_para.text = exp.title
            title_para.add_run('\n\n')
            link_run = title_para.add_run(exp.github)
            link_run.font.size = Pt(9)
            link_run.font.color.rgb = RGBColor(0, 0, 255)
            link_run.underline = True

            qr_img_data = create_qr_code(exp.github, size=150)
            qr_filename = f"/tmp/qr_{idx}.png"

            with open(qr_filename, 'wb') as f:
                f.write(qr_img_data.read())

            qr_images.append(qr_filename)

            cells[3].text = ''
            paragraph = cells[3].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(qr_filename, width=Inches(0.75))

            cells[4].text = ''
            cells[5].text = ''

        doc.add_paragraph()
        doc.add_paragraph()

        confirmation = doc.add_paragraph()
        conf_run = confirmation.add_run(
            'I confirm that the experiments and GitHub links provided are entirely my own work.'
        )
        conf_run.bold = True

        doc.add_paragraph()

        details_table = doc.add_table(rows=2, cols=2)
        details_table.autofit = False

        name_cell = details_table.rows[0].cells[0]
        reg_cell = details_table.rows[0].cells[1]

        name_para = name_cell.paragraphs[0]
        name_para.add_run(f'Name: {data.student_name}').font.size = Pt(11)

        reg_para = reg_cell.paragraphs[0]
        reg_para.add_run(f'Register Number: {data.register_number}').font.size = Pt(11)
        reg_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        output_filename = f"/tmp/{data.register_number}_Lab_Record.docx"
        doc.save(output_filename)

        response = FileResponse(
            path=output_filename,
            filename=output_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        for qr_file in qr_images:
            if os.path.exists(qr_file):
                os.remove(qr_file)

        return response
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
